import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_preprint_docx(output_path):
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Anchor Runtime v5.0.4', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Deterministic Governance Infrastructure for Institutional Agentic Systems')
    subtitle_run = subtitle.runs[0]
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author_info = doc.add_paragraph('Tanishq Dasari\nAnimusLab Research\ncontact@anchorgovernance.tech\nhttps://anchorgovernance.tech')
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run('_' * 40)

    # Classification Table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Classification'
    cells[1].text = 'Institutional Infrastructure Preprint'
    cells = table.rows[1].cells
    cells[0].text = 'Version'
    cells[1].text = '5.0.4 - DETERMINISTIC'
    cells = table.rows[2].cells
    cells[0].text = 'Domain'
    cells[1].text = 'Sovereign Execution, AI Runtime Governance'

    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', 1)
    doc.add_paragraph(
        "Anchor Runtime v5.0.4 is not a governance assistant, monitoring platform, compliance wrapper, or policy layer. "
        "It is a constitutional execution infrastructure for sovereign agentic systems."
    )
    doc.add_paragraph(
        "Modern AI systems are transitioning from deterministic software toward autonomous operational entities. "
        "Existing governance models remain observational. Anchor Runtime transforms institutional policy into "
        "executable constitutional infrastructure capable of deterministic authority propagation, capability enforcement, "
        "and evidentiary lineage continuity. Governance becomes an infrastructure primitive, not an afterthought."
    )

    # 1. Collapse of Observational Governance
    doc.add_heading('1. The Collapse of Observational Governance', 1)
    doc.add_paragraph(
        "Most AI governance systems today are architecturally reactive. They observe behavior after execution. "
        "The problem is no longer model safety; the problem is runtime sovereignty. Traditional architectures fail "
        "because they rely on mutable RBAC and lack deterministic authority propagation."
    )

    # The 4 Verbs
    doc.add_heading('2. The Four Constitutional Verbs', 1)
    doc.add_paragraph(
        "Anchor Runtime v5.0.4 operates through four deterministic verbs that define its cognitive execution cycle:"
    )
    
    verbs = {
        "DECLARE": "Formalization of operational intent via .anchor manifests. Rule evaluation is mandatory before execution.",
        "ENFORCE": "Runtime interception and authority gating. Anchor SDK patches execution to enable 'hard-exit' on violation.",
        "DETECT": "Parallel verification of architectural and semantic drift using the Diamond Cage sandbox.",
        "PROVE": "Generation of the Decision Audit Chain (DAC)—an append-only, cryptographically sealed record."
    }
    
    for verb, desc in verbs.items():
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{verb}: ")
        run.bold = True
        p.add_run(desc)

    # 3. Architecture
    doc.add_heading('3. Anchor Runtime Architecture', 1)
    doc.add_paragraph("[Diagram 1: Sovereign Infrastructure Stack - Horizontal Layering]")
    
    doc.add_heading('3.1 Operational Gateway Architecture', 2)
    gw_table = doc.add_table(rows=1, cols=4)
    gw_table.style = 'Table Grid'
    hdr_cells = gw_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Function'
    hdr_cells[2].text = 'Persistence'
    hdr_cells[3].text = 'Visibility'
    
    gw_data = [
        ["Sovereign Spoke", "Local Audit Ingress", "SQLite / Persistent", "FULL (Enterprise Bound)"],
        ["Relay Gateway", "Metadata Transformation", "Transient / Buffer", "ENCRYPTED (HMAC-SHA256)"],
        ["Federated Hub", "Governance Coordination", "Global Ledger / Postgres", "METADATA-ONLY"],
        ["Oversight Node", "Regulator Transparency", "Read-Only View", "JURISDICTION-SCOPED"]
    ]
    
    for row in gw_data:
        row_cells = gw_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # 3.2 Global Governance Domains (The 9 Primitives)
    doc.add_heading('3.2 Global Governance Domains', 2)
    doc.add_paragraph("Anchor v5.0.4 governs the mesh via 9 fundamental semantic domains:")
    
    domains = {
        "SEC (Security)": "Verification of input/output for malicious injections and exfiltration.",
        "ETH (Ethics)": "Aho-Corasick proxy scanning for behavioral alignment.",
        "PRV (Privacy)": "PII and sensitive data extraction gating.",
        "ALN (Alignment)": "Enforcing declared intent against execution drift.",
        "AGT (Agency)": "Monitoring recursive tool invocation and delegated authority.",
        "LEG (Legal)": "Direct mapping to regulatory dialetics.",
        "OPS (Operations)": "Service-level and configuration integrity.",
        "SUP (Supply)": "Third-party dependency and provenance tracking.",
        "SHR (Sharing)": "Cross-hub data synchronization boundaries."
    }
    
    for dom, ddesc in domains.items():
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{dom}: ")
        run.bold = True
        p.add_run(ddesc)

    doc.add_heading('3.3 Institutional Oversight (The 6 Dialects)', 2)
    doc.add_paragraph("The runtime implements native translation for major regulatory frameworks:")
    regulators = ["RBI (Reserve Bank of India)", "EU (EU AI Act)", "SEC (US Securities & Exchange)", "SEBI (India Securities)", "CFPB (US Consumer Finance)", "FCA (UK Financial Conduct)"]
    for reg in regulators:
        doc.add_paragraph(reg, style='List Bullet')

    doc.add_heading('3.4 Key API Endpoints (v5.0.4)', 2)
    ep_table = doc.add_table(rows=1, cols=4)
    ep_table.style = 'Table Grid'
    hdr_cells = ep_table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Endpoint'
    hdr_cells[2].text = 'Method'
    hdr_cells[3].text = 'Purpose'
    
    ep_data = [
        ["Identity", "/api/auth/identify-first", "POST", "Initial Clearance ID lookup"],
        ["Identity", "/api/auth/me", "GET", "Profile & Institutional context"],
        ["Forensics", "/api/forensic/request", "POST", "Auditor permission escalation"],
        ["Forensics", "/api/forensic/replay/{id}", "GET", "Deterministic chain reconstruction"],
        ["Audit", "/api/ledger", "GET/POST", "Ingress/Query for governance events"],
        ["Relay", "/ws/spoke", "WSS", "Persistent Spoke-Hub handshake"],
        ["Integrity", "/api/audit/{id}/verify", "GET", "Cryptographic chain validation"]
    ]
    
    for row in ep_data:
        row_cells = ep_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # 4. Forensic Replay
    doc.add_heading('4. Forensic Reconstruction', 1)
    doc.add_paragraph("[Diagram 2: Evidence Continuity Chain - JWT -> Exec -> DAC -> Replay]")
    doc.add_paragraph(
        "Anchor treats replay as a sovereign forensic requirement. Raw payloads NEVER leave the enterprise perimeter via REST. "
        "They only travel over the brokered WebSocket relay, AES-256-GCM encrypted, on explicit authorization. "
        "Forensic Pull requests are themselves logged as governance events."
    )

    # 5. Drift Detection
    doc.add_heading('5. Differential Behavioral Verification (Diamond Cage)', 1)
    doc.add_paragraph(
        "The Diamond Cage is a WASM-isolated sandbox that performs side-by-side comparison of original and patched codebases. "
        "It detects architectural drift where logic expands beyond its original semantic intent."
    )

    # 6. Strategic Implications
    doc.add_heading('6. Strategic Implications', 1)
    doc.add_paragraph(
        "For institutions like J.P. Morgan or BlackRock, the game is not about 'buying AI.' "
        "It is about architecting the systems those institutions depend on. Anchor is the black-box recorder and "
        "governance kernel for autonomous capital."
    )

    # Footer
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = "Institutional Integrity Guaranteed by Anchor Governance v5.0.4 | AnimusLab Infrastructure"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)

if __name__ == "__main__":
    output_docx = "Anchor_v5_0_4_Institutional_Preprint.docx"
    create_preprint_docx(output_docx)
    print(f"Document created successfully: {output_docx}")
